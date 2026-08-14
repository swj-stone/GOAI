from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter

from campusmatch.services.document_service import DocumentError, convert_document


def make_docx() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_heading("课程经历", level=1)
    document.add_paragraph("协调小组完成课堂展示。")
    document.save(buffer)
    return buffer.getvalue()


def make_blank_pdf() -> bytes:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buffer)
    return buffer.getvalue()


def test_markdown_upload_keeps_stable_text() -> None:
    result = convert_document("resume.md", "# 经历\r\n\r\n使用 Excel 整理名单。\r\n".encode())

    assert result.source_format == "md"
    assert result.markdown == "# 经历\n\n使用 Excel 整理名单。"
    assert result.char_count == len(result.markdown)


def test_docx_upload_becomes_readable_markdown() -> None:
    result = convert_document("resume.docx", make_docx())

    assert result.source_format == "docx"
    assert "# 课程经历" in result.markdown
    assert "协调小组完成课堂展示。" in result.markdown


def test_scanned_pdf_requests_text_instead_of_silently_succeeding() -> None:
    with pytest.raises(DocumentError) as error:
        convert_document("scan.pdf", make_blank_pdf())

    assert error.value.code == "PDF_TEXT_NOT_FOUND"


def test_unsupported_or_oversized_upload_is_rejected() -> None:
    with pytest.raises(DocumentError) as unsupported:
        convert_document("resume.exe", b"not a document")
    with pytest.raises(DocumentError) as oversized:
        convert_document("resume.txt", b"a" * (5 * 1024 * 1024 + 1))

    assert unsupported.value.code == "DOCUMENT_TYPE_UNSUPPORTED"
    assert oversized.value.code == "DOCUMENT_TOO_LARGE"
