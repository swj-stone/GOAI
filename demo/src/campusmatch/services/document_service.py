from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from campusmatch.contracts import DocumentConversion


MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
SUPPORTED_DOCUMENT_TYPES = {".md", ".txt", ".docx", ".pdf"}


class DocumentError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def normalize_markdown(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    while lines and not lines[-1]:
        lines.pop()

    if not any(line.strip() for line in lines):
        raise DocumentError("DOCUMENT_EMPTY", "材料为空，请上传文件或粘贴文字。")

    return "\n".join(lines)


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError(
        "DOCUMENT_ENCODING_UNSUPPORTED",
        "文本编码无法识别，请另存为 UTF-8 后重试。",
    )


def docx_to_markdown(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise DocumentError("DOCUMENT_PARSE_FAILED", "DOCX 文件无法读取或已损坏。") from error

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading "):
            try:
                level = min(max(int(style.removeprefix("Heading ")), 1), 6)
            except ValueError:
                level = 2
            blocks.append(f"{'#' * level} {text}")
        elif style.startswith("List"):
            blocks.append(f"- {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows or not any(any(cell for cell in row) for row in rows):
            continue
        width = max(len(row) for row in rows)
        padded = [row + [""] * (width - len(row)) for row in rows]
        blocks.append("| " + " | ".join(padded[0]) + " |")
        blocks.append("| " + " | ".join(["---"] * width) + " |")
        blocks.extend("| " + " | ".join(row) + " |" for row in padded[1:])

    return normalize_markdown("\n\n".join(blocks))


def pdf_to_markdown(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        pages = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            pages.append(normalize_markdown(extracted) if extracted.strip() else "")
    except Exception as error:
        raise DocumentError("DOCUMENT_PARSE_FAILED", "PDF 文件无法读取或已损坏。") from error

    pages = [page for page in pages if page.strip()]
    if not pages:
        raise DocumentError(
            "PDF_TEXT_NOT_FOUND",
            "PDF 中没有可提取文字；如果是扫描件，请先进行 OCR 或粘贴文字。",
        )
    if len(pages) == 1:
        return pages[0]
    return normalize_markdown(
        "\n\n".join(f"## 第 {index} 页\n\n{page}" for index, page in enumerate(pages, 1))
    )


def convert_document(filename: str, content: bytes) -> DocumentConversion:
    safe_name = Path(filename or "").name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_TYPES:
        raise DocumentError(
            "DOCUMENT_TYPE_UNSUPPORTED",
            "仅支持 Markdown、TXT、DOCX 和可提取文字的 PDF。",
        )
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentError(
            "DOCUMENT_TOO_LARGE",
            "单个文件不能超过 5 MB，请删除无关图片或拆分材料。",
        )
    if not content:
        raise DocumentError("DOCUMENT_EMPTY", "文件为空，请重新选择材料。")

    if suffix in {".md", ".txt"}:
        markdown = normalize_markdown(decode_text(content))
    elif suffix == ".docx":
        markdown = docx_to_markdown(content)
    else:
        markdown = pdf_to_markdown(content)

    return DocumentConversion(
        filename=safe_name,
        source_format=suffix.removeprefix("."),
        markdown=markdown,
        char_count=len(markdown),
    )
